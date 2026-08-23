#!/usr/bin/env python3
"""Build a fixed-format short or long WeChat PA manuscript."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


REQUIRED_FIELDS = (
    "english_title",
    "chinese_title",
    "authors",
    "chinese_abstract",
    "english_abstract",
    "citation",
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--cover", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--long-summary", type=Path)
    return parser.parse_args()


def load_input(path: Path) -> dict[str, str]:
    raw = json.loads(path.read_text(encoding="utf-8"))
    missing = [field for field in REQUIRED_FIELDS if not str(raw.get(field, "")).strip()]
    if missing:
        raise ValueError(f"Missing non-empty fields: {', '.join(missing)}")

    document_type = str(raw.get("document_type", "short")).strip().lower()
    if document_type not in {"short", "long"}:
        raise ValueError("document_type must be 'short' or 'long'")

    data = {field: str(raw[field]).strip() for field in REQUIRED_FIELDS}
    data["document_type"] = document_type
    return data


def load_long_summary_blocks(path: Path) -> list[str]:
    if not path.is_file():
        raise FileNotFoundError(f"Long-summary Markdown not found: {path}")

    text = path.read_text(encoding="utf-8").replace("\r\n", "\n").strip()
    label_match = re.search(r"(?m)^长摘要：\s*$", text)
    if label_match:
        text = text[label_match.end() :].strip()
    if not text:
        raise ValueError("Long-summary Markdown has no content")

    blocks = [
        " ".join(line.strip() for line in block.splitlines() if line.strip()).strip()
        for block in re.split(r"\n\s*\n", text)
    ]
    blocks = [block for block in blocks if block]
    if not blocks:
        raise ValueError("Long-summary Markdown has no non-empty blocks")
    return blocks


def set_run_font(run) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "Songti SC")


def set_paragraph_format(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    for run in paragraph.runs:
        set_run_font(run)


def add_text(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    set_run_font(paragraph.add_run(text))
    set_paragraph_format(paragraph)


def add_blank(document: Document) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_format(paragraph)


def main() -> None:
    args = parse_args()
    if args.output.exists():
        raise FileExistsError(f"Refusing to overwrite: {args.output}")
    if not args.cover.is_file():
        raise FileNotFoundError(f"Cover not found: {args.cover}")

    data = load_input(args.input)
    long_blocks: list[str] = []
    if data["document_type"] == "long":
        if args.long_summary is None:
            raise ValueError("--long-summary is required for document_type=long")
        long_blocks = load_long_summary_blocks(args.long_summary)
    elif args.long_summary is not None:
        raise ValueError("--long-summary requires document_type=long")

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1

    cover_paragraph = document.add_paragraph()
    cover_paragraph.add_run().add_picture(str(args.cover), width=Inches(6.5))
    set_paragraph_format(cover_paragraph)
    add_blank(document)

    blocks = (
        data["english_title"],
        data["chinese_title"],
        f'今天为大家带来的是{data["authors"]}的研究：《{data["chinese_title"]}》。',
        "摘要：",
        data["chinese_abstract"],
        data["english_abstract"],
        "Citation:",
        data["citation"],
    )
    for index, block in enumerate(blocks):
        add_text(document, block)
        if index != len(blocks) - 1:
            add_blank(document)

    if long_blocks:
        add_blank(document)
        add_text(document, "长摘要：")
        add_blank(document)
        for index, block in enumerate(long_blocks):
            add_text(document, block)
            if index != len(long_blocks) - 1:
                add_blank(document)

    document.core_properties.title = data["english_title"]
    document.core_properties.subject = (
        "Public Administration WeChat long summary"
        if long_blocks
        else "Public Administration WeChat short summary"
    )
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.save(args.output)


if __name__ == "__main__":
    main()
