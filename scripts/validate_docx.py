#!/usr/bin/env python3
"""Validate a fixed-format short WeChat PA manuscript."""

from __future__ import annotations

import argparse
import hashlib
import json
from pathlib import Path
import sys
from zipfile import ZipFile
from xml.etree import ElementTree

from docx import Document
from docx.oxml.ns import qn


REQUIRED_FIELDS = (
    "english_title",
    "chinese_title",
    "authors",
    "chinese_abstract",
    "english_abstract",
    "citation",
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", required=True, type=Path)
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--cover", required=True, type=Path)
    return parser.parse_args()


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def validate_docx(docx_path: Path, input_path: Path, cover_path: Path) -> dict:
    data = json.loads(input_path.read_text(encoding="utf-8"))
    expected = {field: str(data.get(field, "")).strip() for field in REQUIRED_FIELDS}
    document_type = str(data.get("document_type", "short")).strip().lower()
    expected_additional_text = [
        str(value).strip()
        for value in data.get("expected_additional_text", [])
        if str(value).strip()
    ]
    document = Document(docx_path)
    paragraph_texts = [paragraph.text for paragraph in document.paragraphs]
    full_text = "\n".join(paragraph_texts)
    intro = f'今天为大家带来的是{expected["authors"]}的研究：《{expected["chinese_title"]}》。'

    checks: dict[str, bool] = {
        "input_fields_nonempty": all(expected.values()),
        "english_title": expected["english_title"] in full_text,
        "chinese_title": expected["chinese_title"] in full_text,
        "intro": intro in full_text,
        "abstract_label": "摘要：" in paragraph_texts,
        "chinese_abstract": expected["chinese_abstract"] in full_text,
        "english_abstract": expected["english_abstract"] in full_text,
        "citation_label": "Citation:" in paragraph_texts,
        "citation": expected["citation"] in full_text,
        "document_type": document_type in {"short", "long"},
        "paragraph_count": (
            len(paragraph_texts) == 17
            if document_type == "short"
            else len(paragraph_texts) > 17
        ),
        "nonempty_paragraph_count": (
            sum(bool(text) for text in paragraph_texts) == 8
            if document_type == "short"
            else sum(bool(text) for text in paragraph_texts) > 8
        ),
        "long_summary_label": (
            True if document_type == "short" else "长摘要：" in paragraph_texts
        ),
        "additional_text": all(
            value in full_text for value in expected_additional_text
        ),
        "one_inline_image": len(document.inline_shapes) == 1,
        "blank_author": not (document.core_properties.author or "").strip(),
        "blank_last_modified_by": not (
            document.core_properties.last_modified_by or ""
        ).strip(),
    }

    text_runs = [
        run
        for paragraph in document.paragraphs
        for run in paragraph.runs
        if run.text
    ]
    checks["western_font"] = bool(text_runs) and all(
        run._element.get_or_add_rPr().rFonts is not None
        and run._element.get_or_add_rPr().rFonts.get(qn("w:ascii"))
        == "Times New Roman"
        and run._element.get_or_add_rPr().rFonts.get(qn("w:hAnsi"))
        == "Times New Roman"
        for run in text_runs
    )
    checks["east_asian_font"] = bool(text_runs) and all(
        run._element.get_or_add_rPr().rFonts is not None
        and run._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia"))
        == "Songti SC"
        for run in text_runs
    )

    with ZipFile(docx_path) as archive:
        names = archive.namelist()
        media = [name for name in names if name.startswith("word/media/")]
        document_xml = archive.read("word/document.xml")
        document_root = ElementTree.fromstring(document_xml)
        checks["one_media_part"] = len(media) == 1
        checks["cover_hash"] = (
            len(media) == 1
            and sha256_bytes(archive.read(media[0]))
            == sha256_bytes(cover_path.read_bytes())
        )
        checks["no_comments"] = not any(
            "/comments" in name or name.endswith("people.xml")
            for name in names
        )
        checks["no_tracked_changes"] = all(
            document_root.find(f".//{qn(tag)}") is None
            for tag in ("w:ins", "w:del", "w:moveFrom", "w:moveTo")
        )

    return {"ok": all(checks.values()), "checks": checks}


def main() -> None:
    args = parse_args()
    result = validate_docx(args.docx, args.input, args.cover)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    if not result["ok"]:
        sys.exit(1)


if __name__ == "__main__":
    main()
